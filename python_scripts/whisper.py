import subprocess
import re
import chardet
import os
import deepl
from pathlib import Path
from docx import Document
from google_gemini_ai import process_file
from argparse import ArgumentParser

FILE_EXTENSIONS = ['.mkv', '.mp4', '.mp3', '.flac', '.m4a', '.ogg', '.aac', '.opus', '.wmv', '.ts', '.flv', '.avi']


def whisper_logic(file: Path, model: str, language: str):
    if file.suffix.lower() not in FILE_EXTENSIONS:
        return

    subtitle_file = file.with_suffix('.srt')

    if subtitle_file.exists():
        print(f"Subtitle for {file.stem} already exists. Skipping...")
        return

    print(f"Now transcribing: {file.name}")

    subprocess.run(['whisper', '--fp16', 'False', '--output_format', 'srt', '--output_dir', str(file.parent), '--model', model, '--language', language, str(file)])

    remove_subtitle_duplication(subtitle_file)

    if language == "Japanese":
    #     # new_file = process_file(input_file=subtitle_file, instructions="Translate to English whilst retaining SRT formatting without removing any lines.")

        translated_text = deepl_translate(subtitle_file.read_text())
        subtitle_file.write_text(translated_text)
        new_file = subtitle_file

        subtitle_file.unlink()
        new_file.rename(subtitle_file.name)
        print(f"Translated {new_file.name} to English.")


def whisp(file: Path):
    whisper_logic(file, "small.en", "English")


def whisper_path(directory: Path):
    for file in directory.glob('*'):
        if file.is_file():
            whisp(file)


def whisper_path_recursive(directory: Path):
    for subdir in directory.rglob('*'):
        if subdir.is_dir():
            whisper_path(subdir)
    
    whisper_path(directory)


def whisper_japanese(file: Path):
    whisper_logic(file, "small", "Japanese")


def whisper_path_japanese(directory: Path):
    for file in directory.glob('*'):
        if file.is_file():
            whisper_japanese(file)


def remove_subtitle_duplication(file: Path):
    old_text = r'(\d+\r?\n\d+.*?\r?\n(.*?))(?:\r?\n)+(?:\d+\r?\n\d+.*?\r?\n\2(?:\r?\n)+)+'  
    new_text = r'\1\n\n'

    if file.exists():
        with open(file, 'r', encoding='utf-8') as f:
            content = f.read()

        new_content = re.sub(old_text, new_text.strip(), content)

        with open(file, 'w', encoding='utf-8') as f:
            f.write(new_content)
    else:
        print(f"{file} not found.")


def deepl_translate(input_text):
    translated_text = deepl.Translator(os.getenv("DEEPL_API_KEY")).translate_text(input_text, target_lang='EN-US').text

    return translated_text


def srt_to_word(input_file: Path):
    with open(input_file, 'rb') as f:
        raw_data = f.read()
        encoding = chardet.detect(raw_data)['encoding']

    with open(input_file, 'r', encoding=encoding) as f:
        doc = Document()
        doc.add_paragraph(f.read())
        
        output_file = input_file.with_suffix('.docx')
        
        doc.save(str(output_file))
        print(f"Output saved to '{output_file}'")


def word_to_srt(input_file: Path):
    doc = Document(str(input_file))
    text = '\n'.join([para.text for para in doc.paragraphs])
    output_file = f'{str(input_file)[:-8]}.srt'
    with open(output_file, 'w', encoding='utf-16') as f:
        f.write(text)
    print(f"Output saved to '{output_file}'")


def main():
    parser = ArgumentParser(description="Process various file types and perform transcription or conversion")

    parser.add_argument("command", choices=["whisper_logic", "whisp", "whisper_path", "whisper_path_recursive", 
                                            "whisper_japanese", "whisper_path_japanese", "srt_to_word", 
                                            "word_to_srt", "rsd"], 
                        help="Command to execute")

    parser.add_argument("path", type=Path, help="Path to the file or directory")
    parser.add_argument("--model", type=str, default="small.en", help="Model for transcription (default: small.en)")
    parser.add_argument("--language", type=str, default="English", help="Language for transcription (default: English)")

    args = parser.parse_args()

    command_map = {
        "whisper_logic": lambda: whisper_logic(args.path, args.model, args.language),
        "whisp": lambda: whisp(args.path),
        "whisper_path": lambda: whisper_path(args.path),
        "whisper_path_recursive": lambda: whisper_path_recursive(args.path),
        "whisper_japanese": lambda: whisper_japanese(args.path),
        "whisper_path_japanese": lambda: whisper_path_japanese(args.path),
        "srt_to_word": lambda: srt_to_word(args.path),
        "word_to_srt": lambda: word_to_srt(args.path),
        "rsd": lambda: remove_subtitle_duplication(args.path),
    }

    command_map.get(args.command, lambda: print("Invalid command entered."))()


if __name__ == "__main__":
    main()