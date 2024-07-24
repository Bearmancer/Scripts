import deepl, sys, os
from docx import Document
auth_key = open("DeepL API Key.txt", 'r').read()

auth_key = "" 
translator = deepl.Translator(auth_key)

if len(sys.argv) != 2:
    sys.exit(1)

doc = Document(sys.argv[1])

text = []

for para in doc.paragraphs:
    text.append(para.text)

full_text = '\n'.join(text)

result = translator.translate_text(full_text, target_lang="EN-GB")

translated_doc = Document()
translated_doc.add_paragraph(result.text)

base, ext = os.path.splitext(sys.argv[1])
new_file_path = f"{base} - Translated{ext}"

translated_doc.save(new_file_path)