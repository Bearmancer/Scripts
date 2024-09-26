import subprocess, sys, re, chardet
from pathlib import Path
from docx import Document

file_extensions = ['.mkv', '.mp4', '.mp3', '.flac', '.m4a', '.ogg', '.opus', '.wmv', '.ts', '.flv', '.avi']

def whisper_logic(file: Path, model, language):
    file = Path(file)

    if file.suffix not in file_extensions:
        print(f"{file}'s extension ({file.suffix}) is incompatibile. Terminating.")
        return

    subtitle_file = file.with_suffix('.srt')

    if subtitle_file.exists():
        print(f"Subtitle for {file.stem} already exists. Skipping...")
        return

    print(f"Now transcribing: {file.name}")

    subprocess.run(['whisper', '--fp16', 'False', '--output_format', 'srt', '--model', model, '--language', language, str(file)])

    remove_subtitle_duplication(subtitle_file)

    if language == "Japanese":
        srt_to_word(subtitle_file)

def whisp(file):
    whisper_logic(file, "small.en", "English")

def whisper_path(directory):
    for file in directory.glob('*'):
        if file.is_file():
            whisp(file)

def whisper_path_recursive(directory):
    for subdir in directory.rglob('*'):
        if subdir.is_dir():
            whisper_path(subdir)
    whisper_path(directory)

def whisper_japanese(file):
    whisper_logic(file, "small", "Japanese")

def whisper_path_japanese(directory):
    for file in directory.glob('*'):
        if file.is_file():
            whisper_japanese(file)

def remove_subtitle_duplication(file):
    old_text = r'(\d+\r?\n\d+.*?\r?\n(.*?))(?:\r?\n)+(?:\d+\r?\n\d+.*?\r?\n\2(?:\r?\n)+)+'
    new_text = r'\1\n\n'

    if file.exists():
        with open(file, 'r', encoding='utf-16') as f:
            content = f.read()

        new_content = re.sub(old_text, new_text.strip(), content)

        with open(file, 'w', encoding='utf-16') as f:
            f.write(new_content)
    else:
        print(f"{file} not found.")

def srt_to_word(input_file):
    with open(input_file, 'rb') as f:
        raw_data = f.read()
        encoding = chardet.detect(raw_data)['encoding']

    with open(input_file, 'r', encoding=encoding) as f:
        doc = Document()
        doc.add_paragraph(f.read())
        output_file = input_file.replace('.srt', '.docx')
        doc.save(output_file)
        print(f"Output saved to '{output_file}")

def word_to_srt(input_file):
    doc = Document(input_file)
    text = '\n'.join([para.text for para in doc.paragraphs])
    output_file = f'{str(input_file)[:-8]}.srt'
    with open(output_file, 'w', encoding='utf-16') as f:
        f.write(text)
    print(f"Output saved to '{output_file}'")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Invalid input entered.")
        exit()

    command = sys.argv[1]
    path = Path(sys.argv[2])

    if command == "WhisperLogic":
        whisper_logic(path, sys.argv[3], sys.argv[4])
    elif command == "Whisp":
        whisp(path)
    elif command == "WhisperPath":
        whisper_path(path)
    elif command == "WhisperPathRecursive":
        whisper_path_recursive(path)
    elif command == "WhisperJapanese":
        whisper_japanese(path)
    elif command == "WhisperPathJapanese":
        whisper_path_japanese(path)
    elif command == "SRTtoWord":
        srt_to_word(path)
    elif command == "WordToSRT":
        word_to_srt(path)
    elif command == "rsd":
        remove_subtitle_duplication(path)
    else:
        print("Invalid command entered.")