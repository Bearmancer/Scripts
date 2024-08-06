from docx import Document
import sys, chardet, os

def srt_to_docx(input_file):
    with open(input_file, 'rb') as f:
        raw_data = f.read()
        encoding = chardet.detect(raw_data)['encoding']

    with open(input_file, 'r', encoding=encoding) as f:
        doc = Document()
        doc.add_paragraph(f.read())
        output_file = input_file.replace('.srt', '.docx')
        doc.save(output_file)
        print(f"Output saved to '{output_file}'")

def docx_to_srt(input_file):
    doc = Document(input_file)
    text = '\n'.join([para.text for para in doc.paragraphs])
    output_file = input_file.replace('.docx', '.srt')
    with open(output_file, 'w', encoding='utf-16') as f:
        f.write(text)
    print(f"Output saved to '{output_file}'")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Invalid input")
        sys.exit(1)

    input_file = sys.argv[1]
    mode = sys.argv[2]

    if not os.path.isfile(input_file):
        print(f"Error: File '{input_file}' not found.")
        sys.exit(1)

    if mode == 'srt':
        srt_to_docx(input_file)
    elif mode == 'docx':
        docx_to_srt(input_file)
    else:
        print("Invalid mode. Use 'srt' or 'docx'.")
        sys.exit(1)